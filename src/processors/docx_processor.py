"""
DOCX document processor for extracting structured content, metadata, and creating content chunks.

Uses python-docx for DOCX parsing with support for text extraction, headings,
tables, document structure, and formatting metadata preservation.
"""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx>=0.8.11")

from .base import DocumentProcessor, ProcessingError, ContentExtractionError, FileValidationError
from ..models import (
    DocumentMetadata, ContentChunk, ValidationResult, ContentType, SourceLocation
)
from ..config import ProcessingConfig


class DOCXProcessor(DocumentProcessor):
    """
    DOCX document processor that extracts structured content and metadata.
    
    Features:
    - Text extraction with structure preservation
    - Heading hierarchy detection and preservation
    - Table content extraction with structure
    - Document metadata extraction
    - Paragraph-level source location tracking
    - Configurable content chunking with structure awareness
    - Formatting metadata preservation
    """
    
    def __init__(self, config: ProcessingConfig):
        super().__init__(config)
        self.supported_formats = ['docx']
    
    def validate_file(self, file_path: str) -> ValidationResult:
        """
        Validate that the file is a readable DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            ValidationResult with validation status and any errors
        """
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                return ValidationResult(
                    is_valid=False,
                    error_message=f"File not found: {file_path}"
                )
            
            # Check file extension
            if path.suffix.lower() != '.docx':
                return ValidationResult(
                    is_valid=False,
                    error_message=f"Invalid file extension. Expected .docx, got {path.suffix}",
                    file_format=path.suffix.lower()
                )
            
            # Check file size
            if self._is_file_too_large(file_path):
                return ValidationResult(
                    is_valid=False,
                    error_message=f"File too large. Maximum size: {self.config.max_file_size_mb}MB",
                    file_size=self._get_file_size(file_path)
                )
            
            # Try to open and read the DOCX
            try:
                doc = Document(file_path)
                
                # Check if document has any content
                if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                    return ValidationResult(
                        is_valid=False,
                        error_message="DOCX document appears to be empty",
                        file_format="docx"
                    )
                
                # Try to access document properties to verify structure
                try:
                    _ = doc.core_properties.title
                except Exception:
                    # Document properties access failed, but document might still be readable
                    pass
            
            except Exception as e:
                return ValidationResult(
                    is_valid=False,
                    error_message=f"Cannot read DOCX file: {str(e)}",
                    file_format="docx"
                )
            
            return ValidationResult(
                is_valid=True,
                file_format="docx",
                file_size=self._get_file_size(file_path)
            )
            
        except Exception as e:
            return ValidationResult(
                is_valid=False,
                error_message=f"Validation error: {str(e)}"
            )
    
    def extract_content(self, file_path: str) -> str:
        """
        Extract structured content from the DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted structured content with headings, paragraphs, and tables
            
        Raises:
            ContentExtractionError: If content extraction fails
        """
        try:
            doc = Document(file_path)
            extracted_content = []
            paragraph_index = 0
            
            # Process document elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    
                    if paragraph:
                        paragraph_content = self._extract_paragraph_content(paragraph, paragraph_index)
                        if paragraph_content.strip():
                            extracted_content.append(paragraph_content)
                            paragraph_index += 1
                
                elif element.tag.endswith('tbl'):  # Table
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        table_content = self._extract_table_content(table, paragraph_index)
                        if table_content.strip():
                            extracted_content.append(table_content)
                            paragraph_index += 1
            
            if not extracted_content:
                raise ContentExtractionError(
                    "No content could be extracted from DOCX document",
                    file_path=file_path
                )
            
            return "\n\n".join(extracted_content)
                
        except FileNotFoundError:
            raise ContentExtractionError(
                f"DOCX file not found: {file_path}",
                file_path=file_path
            )
        except Exception as e:
            raise ContentExtractionError(
                f"Failed to extract content from DOCX: {str(e)}",
                file_path=file_path,
                cause=e
            )
    
    def extract_metadata(self, file_path: str) -> DocumentMetadata:
        """
        Extract metadata from the DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            DocumentMetadata with extracted information
        """
        try:
            doc = Document(file_path)
            
            # Get basic file information
            file_size = self._get_file_size(file_path)
            
            # Count paragraphs and tables for page estimation
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)
            
            # Estimate page count (rough approximation)
            # Assume ~25 paragraphs per page or ~2 tables per page
            estimated_pages = max(1, (paragraph_count // 25) + (table_count // 2))
            
            # Extract document properties
            props = doc.core_properties
            
            # Extract title
            title = props.title if props.title else Path(file_path).stem
            
            # Extract author
            author = props.author if props.author else None
            
            # Extract creation date
            creation_date = None
            if props.created:
                creation_date = props.created.strftime("%Y-%m-%d")
            
            # Extract language (if available)
            language = None
            try:
                # Try to get language from document settings
                settings = doc.settings
                if hasattr(settings, 'language'):
                    language = settings.language
            except Exception:
                pass
            
            # Get format version info
            format_version = "DOCX (Office Open XML)"
            
            return DocumentMetadata(
                title=title,
                author=author,
                creation_date=creation_date,
                file_size=file_size,
                page_count=estimated_pages,
                format_version=format_version,
                language=language,
                confidence_score=1.0
            )
                
        except Exception as e:
            # Return basic metadata if extraction fails
            return DocumentMetadata(
                title=Path(file_path).stem,
                file_size=self._get_file_size(file_path),
                page_count=1,
                confidence_score=0.5
            )
    
    def chunk_content(self, content: str, document_id: str, file_path: str) -> List[ContentChunk]:
        """
        Split DOCX content into chunks with paragraph-level source tracking.
        
        Args:
            content: Processed content to chunk
            document_id: Unique identifier for the document
            file_path: Original file path for source location
            
        Returns:
            List of ContentChunk objects with paragraph information
        """
        chunks = []
        chunk_index = 0
        
        # Split content by structure markers
        sections = self._split_by_structure(content)
        
        for section_info in sections:
            section_content = section_info['content']
            paragraph_index = section_info['paragraph_index']
            section_type = section_info['type']
            
            if not section_content.strip():
                continue
            
            # For headings and short sections, keep as single chunks
            if section_type == 'heading' or len(section_content) <= self.config.chunk_size:
                chunk_id = self._create_chunk_id(document_id, chunk_index)
                
                source_location = SourceLocation(
                    file_path=file_path,
                    paragraph_index=paragraph_index
                )
                
                chunk = ContentChunk(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    content=section_content.strip(),
                    content_type=ContentType.DOCX,
                    source_location=source_location,
                    metadata={
                        'paragraph_index': paragraph_index,
                        'chunk_index': chunk_index,
                        'section_type': section_type,
                        'character_count': len(section_content),
                        'word_count': len(section_content.split()),
                        'heading_level': section_info.get('heading_level')
                    },
                    confidence_score=1.0
                )
                
                chunks.append(chunk)
                chunk_index += 1
            
            else:
                # Split longer sections into chunks
                section_chunks = self._create_text_chunks(
                    section_content,
                    self.config.chunk_size,
                    self.config.chunk_overlap
                )
                
                for chunk_text in section_chunks:
                    if not chunk_text.strip():
                        continue
                    
                    chunk_id = self._create_chunk_id(document_id, chunk_index)
                    
                    source_location = SourceLocation(
                        file_path=file_path,
                        paragraph_index=paragraph_index
                    )
                    
                    chunk = ContentChunk(
                        chunk_id=chunk_id,
                        document_id=document_id,
                        content=chunk_text.strip(),
                        content_type=ContentType.DOCX,
                        source_location=source_location,
                        metadata={
                            'paragraph_index': paragraph_index,
                            'chunk_index': chunk_index,
                            'section_type': section_type,
                            'character_count': len(chunk_text),
                            'word_count': len(chunk_text.split()),
                            'heading_level': section_info.get('heading_level')
                        },
                        confidence_score=1.0
                    )
                    
                    chunks.append(chunk)
                    chunk_index += 1
        
        return chunks
    
    def get_content_type(self) -> ContentType:
        """Get the content type handled by this processor."""
        return ContentType.DOCX
    
    def process_content(self, raw_content: str) -> str:
        """
        Process raw DOCX content by cleaning and normalizing text.
        
        Args:
            raw_content: Raw extracted content
            
        Returns:
            Processed content ready for chunking
        """
        # Remove structure markers for processing
        content = re.sub(r'\[PARAGRAPH \d+\]\n', '', raw_content)
        content = re.sub(r'\[TABLE \d+\]\n', '', raw_content)
        content = re.sub(r'\[HEADING \d+ LEVEL \d+\]\n', '', raw_content)
        
        # Clean up common DOCX extraction artifacts
        content = self._clean_docx_text(content)
        
        return content.strip()
    
    def _extract_paragraph_content(self, paragraph, paragraph_index: int) -> str:
        """
        Extract content from a paragraph with formatting information.
        
        Args:
            paragraph: python-docx Paragraph object
            paragraph_index: Index of the paragraph in the document
            
        Returns:
            Formatted paragraph content with structure markers
        """
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # Determine if this is a heading
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        is_heading = 'heading' in style_name
        
        if is_heading:
            # Extract heading level
            heading_level = 1
            if 'heading' in style_name:
                try:
                    level_match = re.search(r'heading\s*(\d+)', style_name)
                    if level_match:
                        heading_level = int(level_match.group(1))
                except (ValueError, AttributeError):
                    pass
            
            return f"[HEADING {paragraph_index} LEVEL {heading_level}]\n{text}"
        else:
            return f"[PARAGRAPH {paragraph_index}]\n{text}"
    
    def _extract_table_content(self, table, paragraph_index: int) -> str:
        """
        Extract content from a table with structure preservation.
        
        Args:
            table: python-docx Table object
            paragraph_index: Index position in the document
            
        Returns:
            Formatted table content with structure markers
        """
        table_content = []
        table_content.append(f"[TABLE {paragraph_index}]")
        
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Replace newlines in cells with spaces
                cell_text = re.sub(r'\s+', ' ', cell_text)
                row_cells.append(cell_text)
            
            # Join cells with tab separator
            row_text = "\t".join(row_cells)
            if row_text.strip():
                table_content.append(row_text)
        
        return "\n".join(table_content)
    
    def _split_by_structure(self, content: str) -> List[Dict[str, Any]]:
        """
        Split content by structure markers and return structured sections.
        
        Args:
            content: Content with structure markers
            
        Returns:
            List of dictionaries with section information
        """
        sections = []
        lines = content.split('\n')
        current_section = []
        current_paragraph_index = 0
        current_type = 'paragraph'
        current_heading_level = None
        
        for line in lines:
            # Check for structure markers
            paragraph_match = re.match(r'\[PARAGRAPH (\d+)\]', line)
            table_match = re.match(r'\[TABLE (\d+)\]', line)
            heading_match = re.match(r'\[HEADING (\d+) LEVEL (\d+)\]', line)
            
            if paragraph_match or table_match or heading_match:
                # Save previous section
                if current_section:
                    sections.append({
                        'content': '\n'.join(current_section),
                        'paragraph_index': current_paragraph_index,
                        'type': current_type,
                        'heading_level': current_heading_level
                    })
                
                # Start new section
                if paragraph_match:
                    current_paragraph_index = int(paragraph_match.group(1))
                    current_type = 'paragraph'
                    current_heading_level = None
                elif table_match:
                    current_paragraph_index = int(table_match.group(1))
                    current_type = 'table'
                    current_heading_level = None
                elif heading_match:
                    current_paragraph_index = int(heading_match.group(1))
                    current_type = 'heading'
                    current_heading_level = int(heading_match.group(2))
                
                current_section = []
            else:
                current_section.append(line)
        
        # Save last section
        if current_section:
            sections.append({
                'content': '\n'.join(current_section),
                'paragraph_index': current_paragraph_index,
                'type': current_type,
                'heading_level': current_heading_level
            })
        
        return sections
    
    def _create_text_chunks(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        Create overlapping text chunks from content.
        
        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # If this is not the last chunk, try to break at word boundary
            if end < len(text):
                # Look for the last space within the chunk
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap
            if start <= 0:
                start = end
        
        return chunks
    
    def _clean_docx_text(self, text: str) -> str:
        """
        Clean common DOCX text extraction artifacts.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common spacing issues
        text = re.sub(r'\s+([\.!?])', r'\1', text)  # Remove space before punctuation
        text = re.sub(r'([\.!?])\s*([A-Z])', r'\1 \2', text)  # Ensure space after sentence end
        
        # Clean up table formatting artifacts
        text = re.sub(r'\t+', '\t', text)  # Normalize multiple tabs
        text = re.sub(r'\t', ' | ', text)  # Convert tabs to readable separators
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()